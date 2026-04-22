import os
import re
import pandas as pd
import fitz  # PyMuPDF
from typing import List, Dict, Tuple, Optional
import hashlib
from datetime import datetime

def calculate_content_hash(text: str) -> str:
    """Calculate a hash of the content for duplicate detection."""
    return hashlib.md5(text.encode('utf-8')).hexdigest()

def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extracts all text from a PDF file using PyMuPDF with better error handling.

    Args:
        pdf_path: Path to the PDF file.

    Returns:
        A string with the full content of the PDF.
    """
    try:
        doc = fitz.open(pdf_path)
        text = ""
        
        for page_num in range(len(doc)):
            try:
                page = doc[page_num]
                page_text = page.get_text()
                
                # Clean up the text
                page_text = re.sub(r'\s+', ' ', page_text.strip())
                
                if page_text:
                    text += page_text + "\n\n"
                    
            except Exception as page_error:
                print(f"Warning: Error reading page {page_num + 1}: {page_error}")
                continue
                
        doc.close()
        
        # Final cleanup
        text = text.strip()
        return text
        
    except Exception as e:
        print(f"Error reading file {pdf_path}: {e}")
        return ""

def extract_text_from_txt(txt_path: str) -> str:
    """Extract text from a TXT file with encoding detection."""
    import chardet
    
    try:
        # Try to detect encoding
        with open(txt_path, 'rb') as f:
            raw_data = f.read()
            encoding_info = chardet.detect(raw_data)
            encoding = encoding_info.get('encoding', 'utf-8')
        
        # Read with detected encoding
        with open(txt_path, 'r', encoding=encoding) as f:
            text = f.read()
        
        return text.strip()
        
    except Exception as e:
        print(f"Error reading TXT file {txt_path}: {e}")
        return ""

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        import docx
        doc = docx.Document(docx_path)
        
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        return "\n\n".join(text_parts)
        
    except ImportError:
        raise ImportError("python-docx library is required for DOCX file support")
    except Exception as e:
        print(f"Error reading DOCX file {docx_path}: {e}")
        return ""

def detect_language(text: str) -> str:
    """Simple language detection based on common words."""
    spanish_indicators = ['el', 'la', 'de', 'que', 'y', 'en', 'un', 'es', 'se', 'no', 'te', 'lo', 'le', 'da', 'su', 'por', 'son', 'con', 'para', 'del', 'las', 'una', 'más', 'este', 'esta', 'como', 'pero', 'sus', 'fue', 'ser', 'muy', 'todo', 'era', 'año', 'dos', 'mismo', 'hace', 'tiempo', 'casa', 'vida', 'país', 'póliza', 'seguro', 'cobertura', 'artículo']
    
    text_lower = text.lower()
    spanish_count = sum(1 for word in spanish_indicators if word in text_lower)
    
    if spanish_count > 5:
        return 'es'
    return 'en'

def chunk_document_semantically(text: str, filename: str, chunk_size: int = 2000, overlap: int = 200) -> List[Dict[str, str]]:
    """
    Creates semantic chunks from document text with improved chunking strategy.
    
    Args:
        text: The text content of the document.
        filename: The name of the source file, for reference.
        chunk_size: Target size for semantic chunks.
        overlap: Number of characters to overlap between chunks.

    Returns:
        A list of dictionaries, where each dictionary represents a structured chunk.
    """
    # Clean up the text
    text = re.sub(r'\s+', ' ', text.strip())
    
    if not text or len(text) < 20:
        return []
    
    # For very short texts (less than chunk_size/3), don't split at all
    if len(text) <= chunk_size // 3:
        chunks = []
        content_hash = calculate_content_hash(text)
        language = detect_language(text)
        
        chunks.append({
            'file_name': os.path.basename(filename),
            'article_number': '1',
            'article_title': f"Document: {os.path.basename(filename)}",
            'article_content': text,
            'chunk_type': 'single_document',
            'language': language,
            'content_hash': content_hash,
            'parent_hash': content_hash,
            'chunk_index': 0
        })
        
        return chunks
    
    chunks = []
    content_hash = calculate_content_hash(text)
    language = detect_language(text)
    
    # Strategy 1: Try to split on clear article boundaries if they exist
    article_patterns = [
        r'(?i)(?:artículo|articulo|article)\s*(?:n°?\s*)?\d+',
        r'(?i)(?:sección|seccion|section)\s*(?:n°?\s*)?\d+',
        r'(?i)(?:capítulo|capitulo|chapter)\s*(?:n°?\s*)?\d+',
        r'(?i)(?:cláusula|clausula|clause)\s*(?:n°?\s*)?\d+'
    ]
    
    best_splits = []
    best_pattern = None
    
    for pattern in article_patterns:
        # Use finditer to get positions while preserving the headers
        matches = list(re.finditer(pattern, text))
        if len(matches) > 2:  # Need at least 3 articles for good structure
            # Split text while preserving headers
            splits = []
            start_pos = 0
            
            for match in matches:
                # Add text before this match (if any)
                if match.start() > start_pos:
                    before_text = text[start_pos:match.start()].strip()
                    if before_text and len(before_text) > 50:
                        splits.append(before_text)
                
                # Find the end of this article (start of next match or end of text)
                next_match_start = matches[matches.index(match) + 1].start() if match != matches[-1] else len(text)
                
                # Extract article content including the header
                article_content = text[match.start():next_match_start].strip()
                if article_content and len(article_content) > 50:
                    splits.append(article_content)
                
                start_pos = next_match_start
            
            # Check if this gives us good splits
            if len(splits) >= 3 and all(len(s.strip()) > 50 for s in splits):
                best_splits = splits
                best_pattern = pattern
                break
    
    if best_splits and best_pattern:
        # Found good article structure with preserved headers
        
        for i, chunk in enumerate(best_splits):
            chunk = chunk.strip()
            if not chunk or len(chunk) < 20:
                continue
                
            # Extract article number from the chunk content itself
            article_num_match = re.search(r'(?i)(?:artículo|articulo|article|sección|seccion|section|capítulo|capitulo|chapter|cláusula|clausula|clause)\s*(?:n°?\s*)?(\d+)', chunk)
            article_num = article_num_match.group(1) if article_num_match else str(i + 1)
            
            # Extract title - look for the header and what follows
            title = f"Article {article_num}"
            
            # Try to extract a meaningful title from the article header
            header_patterns = [
                r'(?i)(?:artículo|articulo|article)\s*(?:n°?\s*)?\d+\s*[:.]?\s*([^.\n]{5,100})',
                r'(?i)(?:sección|seccion|section)\s*(?:n°?\s*)?\d+\s*[:.]?\s*([^.\n]{5,100})',
                r'(?i)(?:capítulo|capitulo|chapter)\s*(?:n°?\s*)?\d+\s*[:.]?\s*([^.\n]{5,100})',
                r'(?i)(?:cláusula|clausula|clause)\s*(?:n°?\s*)?\d+\s*[:.]?\s*([^.\n]{5,100})',
                r'^([^.\n]{10,100})'  # Fallback: first line
            ]
            
            for pattern in header_patterns:
                title_match = re.search(pattern, chunk)
                if title_match:
                    title = title_match.group(1).strip()
                    # Clean up title but preserve meaningful content
                    title = re.sub(r'^[:\.\-\s]+', '', title)
                    title = re.sub(r'[:\.\-\s]+$', '', title)
                    if title and len(title) >= 3:
                        break
            
            # Ensure we have a valid title
            if not title or len(title) < 3:
                title = f"Article {article_num}"
            
            # Split large chunks further if needed
            if len(chunk) > chunk_size * 2:
                sub_chunks = split_long_chunk(chunk, chunk_size, overlap)
                for j, sub_chunk in enumerate(sub_chunks):
                    chunks.append({
                        'file_name': os.path.basename(filename),
                        'article_number': f"{article_num}.{j+1}",
                        'article_title': f"{title} (Part {j+1})",
                        'article_content': sub_chunk,
                        'chunk_type': 'article_part',
                        'language': language,
                        'content_hash': calculate_content_hash(sub_chunk),
                        'parent_hash': content_hash,
                        'chunk_index': len(chunks)
                    })
            else:
                chunks.append({
                    'file_name': os.path.basename(filename),
                    'article_number': article_num,
                    'article_title': title,
                    'article_content': chunk,
                    'chunk_type': 'article',
                    'language': language,
                    'content_hash': calculate_content_hash(chunk),
                    'parent_hash': content_hash,
                    'chunk_index': len(chunks)
                })
    
    # Strategy 2: If no clear articles, create semantic chunks by content
    if not chunks or len(chunks) < 2:
        chunks = []  # Reset if we didn't find good articles
        
        # Try to split on paragraph boundaries first
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
        
        if not paragraphs or len(paragraphs) == 1:
            # If no paragraph breaks, split on sentence boundaries
            sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
            paragraphs = sentences
        
        if not paragraphs:  # If still no splits, use the whole text
            paragraphs = [text]
        
        current_chunk = ""
        chunk_num = 1
        
        for para in paragraphs:
            # Check if adding this paragraph would make chunk too big
            potential_length = len(current_chunk) + len(para) + 1
            
            if potential_length > chunk_size and current_chunk and len(current_chunk) > chunk_size // 2:
                # Save current chunk
                title = extract_chunk_title(current_chunk, chunk_num)
                
                chunks.append({
                    'file_name': os.path.basename(filename),
                    'article_number': str(chunk_num),
                    'article_title': title,
                    'article_content': current_chunk.strip(),
                    'chunk_type': 'semantic_chunk',
                    'language': language,
                    'content_hash': calculate_content_hash(current_chunk),
                    'parent_hash': content_hash,
                    'chunk_index': chunk_num - 1
                })
                
                # Start new chunk with overlap if configured
                if overlap > 0 and len(current_chunk) > overlap:
                    current_chunk = current_chunk[-overlap:] + " " + para
                else:
                    current_chunk = para
                chunk_num += 1
            else:
                current_chunk += " " + para if current_chunk else para
        
        # Don't forget the last chunk
        if current_chunk.strip():
            title = extract_chunk_title(current_chunk, chunk_num)
            
            chunks.append({
                'file_name': os.path.basename(filename),
                'article_number': str(chunk_num),
                'article_title': title,
                'article_content': current_chunk.strip(),
                'chunk_type': 'semantic_chunk',
                'language': language,
                'content_hash': calculate_content_hash(current_chunk),
                'parent_hash': content_hash,
                'chunk_index': chunk_num - 1
            })
    
    return chunks

def split_long_chunk(text: str, max_size: int, overlap: int) -> List[str]:
    """Split a long chunk into smaller overlapping pieces."""
    if len(text) <= max_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + max_size
        
        if end >= len(text):
            # Last chunk
            chunks.append(text[start:])
            break
        
        # Try to break at a sentence boundary
        chunk_text = text[start:end]
        last_sentence_end = max(
            chunk_text.rfind('.'),
            chunk_text.rfind('!'),
            chunk_text.rfind('?')
        )
        
        if last_sentence_end > max_size // 2:  # Good sentence break found
            end = start + last_sentence_end + 1
        
        chunks.append(text[start:end])
        start = end - overlap if overlap > 0 else end
    
    return chunks

def extract_chunk_title(text: str, chunk_num: int) -> str:
    """Extract a meaningful title from a chunk of text."""
    # Try different patterns to find a good title
    title_patterns = [
        r'^([A-ZÁÉÍÓÚÑ][^.\n]{10,80})',  # Capitalized start
        r'^([^.\n]{10,100})',            # First line
        r'([^:]{5,80}):',                # Text before colon
        r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\b'  # Title case words
    ]
    
    for pattern in title_patterns:
        match = re.search(pattern, text.strip())
        if match:
            title = match.group(1).strip()
            # Clean up
            title = re.sub(r'^[:\.\-\s]+', '', title)
            title = re.sub(r'[:\.\-\s]+$', '', title)
            if 10 <= len(title) <= 100:
                return title
    
    # Fallback: use first few words
    words = text.strip().split()[:8]
    if words:
        title = ' '.join(words)
        if len(title) > 80:
            title = title[:77] + "..."
        return title
    
    return f"Section {chunk_num}"

def extract_articles_from_text(text: str, filename: str) -> List[Dict[str, str]]:
    """
    Wrapper function to maintain compatibility with existing code.
    Now uses semantic chunking instead of regex-based article extraction.
    """
    return chunk_document_semantically(text, filename)

def process_files_from_directory(directory: str, allowed_extensions: Optional[List[str]] = None) -> pd.DataFrame:
    """
    Processes all supported files in a given directory, extracts content, and returns a DataFrame.

    Args:
        directory: The path to the directory containing files.
        allowed_extensions: List of allowed file extensions. If None, defaults to ['.pdf', '.txt', '.docx']

    Returns:
        A pandas DataFrame with the structured data of all extracted content.
    """
    if allowed_extensions is None:
        allowed_extensions = ['.pdf', '.txt', '.docx']
    
    if not os.path.isdir(directory):
        print(f"Error: The directory '{directory}' was not found.")
        return pd.DataFrame()

    all_articles_data = []
    processed_files = 0
    failed_files = []

    for filename in os.listdir(directory):
        file_ext = os.path.splitext(filename.lower())[1]
        
        if file_ext not in allowed_extensions:
            continue
            
        file_path = os.path.join(directory, filename)
        print(f"Processing file: {filename}")

        try:
            # Extract text based on file type
            text_content = ""
            
            if file_ext == '.pdf':
                text_content = extract_text_from_pdf(file_path)
            elif file_ext == '.txt':
                text_content = extract_text_from_txt(file_path)
            elif file_ext == '.docx':
                text_content = extract_text_from_docx(file_path)
            
            if text_content and len(text_content.strip()) > 20:
                articles = extract_articles_from_text(text_content, filename)
                if articles:
                    all_articles_data.extend(articles)
                    print(f" -> Extracted {len(articles)} chunks from {filename}")
                    processed_files += 1
                else:
                    print(f" -> Warning: No content chunks extracted from {filename}")
            else:
                print(f" -> Warning: No meaningful text extracted from {filename}")
                failed_files.append(filename)
                
        except Exception as e:
            print(f" -> Error processing {filename}: {e}")
            failed_files.append(filename)

    print(f"\nProcessing completed:")
    print(f"  - Successfully processed: {processed_files} files")
    print(f"  - Failed to process: {len(failed_files)} files")
    if failed_files:
        print(f"  - Failed files: {', '.join(failed_files)}")

    if all_articles_data:
        df = pd.DataFrame(all_articles_data)
        print(f"  - Total chunks extracted: {len(df)}")
        return df
    else:
        print("No content was extracted from any files.")
        return pd.DataFrame()

def process_pdfs_from_directory(pdf_directory: str) -> pd.DataFrame:
    """
    Legacy function - now uses the more general process_files_from_directory.
    Processes all PDF files in a given directory, extracts articles, and returns a DataFrame.

    Args:
        pdf_directory: The path to the directory containing PDF files.

    Returns:
        A pandas DataFrame with the structured data of all extracted articles.
    """
    return process_files_from_directory(pdf_directory, ['.pdf'])

def validate_document_content(text: str, min_length: int = 20) -> Tuple[bool, str]:
    """
    Validate if document content is meaningful and processable.
    
    Args:
        text: The text content to validate
        min_length: Minimum length requirement
        
    Returns:
        Tuple of (is_valid, error_message)
    """
    if not text:
        return False, "No text content found"
    
    text = text.strip()
    
    if len(text) < min_length:
        return False, f"Text too short (minimum {min_length} characters required)"
    
    # Check if text is mostly garbled or encrypted
    readable_chars = sum(1 for c in text if c.isalnum() or c.isspace() or c in '.,;:!?()-[]{}"\'-/')
    total_chars = len(text)
    
    if total_chars > 0 and (readable_chars / total_chars) < 0.7:
        return False, "Text appears to be garbled or encoded"
    
    # Check for minimum word count
    words = text.split()
    if len(words) < 5:
        return False, "Text contains too few words"
    
    return True, "Content is valid"

def get_document_statistics(text: str) -> Dict:
    """Get statistics about a document's content."""
    return {
        'character_count': len(text),
        'word_count': len(text.split()),
        'paragraph_count': len([p for p in text.split('\n\n') if p.strip()]),
        'sentence_count': len([s for s in re.split(r'[.!?]+', text) if s.strip()]),
        'language': detect_language(text),
        'content_hash': calculate_content_hash(text)
    }